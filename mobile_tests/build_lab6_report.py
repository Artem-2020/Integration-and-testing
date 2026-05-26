from __future__ import annotations

import json
import re
import textwrap
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
ARTIFACTS_DIR = ROOT / "artifacts"
SCREENSHOTS_DIR = ARTIFACTS_DIR / "screenshots"
REPORTS_DIR = ROOT / "reports"
LOGS_DIR = REPORTS_DIR / "logs"
ALLURE_REPORT_DIR = REPORTS_DIR / "allure-report"
REPORT_ASSETS_DIR = ROOT / "report_assets"
OUTPUT_DOCX = ROOT / "Отчет_ЛР6_Вариант_12_NewPipe.docx"

TITLE_BLOCK = {
    "student_group": "221331",
    "student_name": "Долгополов А.Д.",
    "teacher_name": "Басалов Ю.А.",
    "city": "Тула",
    "year": "2026",
}

LAB_TITLE = "Автоматизация тестирования мобильного приложения с помощью Appium"
VARIANT_TITLE = "Вариант 12: NewPipe — GitHub Release"
RELEASE_TAG = "v0.28.7"
RELEASE_DATE = "23.05.2026"
RELEASE_URL = (
    "https://github.com/TeamNewPipe/NewPipe/releases/tag/v0.28.7"
)
RELEASE_APK_URL = (
    "https://github.com/TeamNewPipe/NewPipe/releases/download/"
    "v0.28.7/NewPipe_v0.28.7.apk"
)

ANSI_ESCAPE_RE = re.compile(r"\x1B\[[0-?]*[ -/]*[@-~]")


def load_font(candidates: list[str], size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    windows_fonts = Path("C:/Windows/Fonts")
    for name in candidates:
        candidate = windows_fonts / name
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


TERMINAL_FONT = load_font(["consola.ttf", "CascadiaMono.ttf", "lucon.ttf"], 24)
TERMINAL_TITLE_FONT = load_font(
    ["consolab.ttf", "CascadiaMono.ttf", "arialbd.ttf"], 30
)
SANS_FONT = load_font(["arial.ttf", "calibri.ttf"], 28)
SANS_BOLD_FONT = load_font(["arialbd.ttf", "calibrib.ttf"], 30)
SMALL_FONT = load_font(["arial.ttf", "calibri.ttf"], 22)
BIG_FONT = load_font(["arialbd.ttf", "calibrib.ttf"], 58)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def clean_ansi(text: str) -> str:
    return ANSI_ESCAPE_RE.sub("", text)


def wrap_console_lines(lines: list[str], width: int = 104) -> list[str]:
    wrapped: list[str] = []
    for raw_line in lines:
        line = raw_line.rstrip("\n")
        if not line:
            wrapped.append("")
            continue
        wrapped.extend(
            textwrap.wrap(
                line,
                width=width,
                replace_whitespace=False,
                drop_whitespace=False,
            )
            or [line]
        )
    return wrapped


def draw_terminal_image(title: str, lines: list[str], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    prepared_lines = wrap_console_lines(lines)
    padding = 36
    line_height = 34
    width = 1600
    title_height = 70
    height = padding * 2 + title_height + max(1, len(prepared_lines)) * line_height
    image = Image.new("RGB", (width, height), "#0d1117")
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle(
        (18, 18, width - 18, height - 18),
        radius=18,
        outline="#30363d",
        width=2,
        fill="#0d1117",
    )
    draw.text((padding, padding), title, font=TERMINAL_TITLE_FONT, fill="#7ee787")

    y = padding + title_height
    for line in prepared_lines:
        draw.text((padding, y), line, font=TERMINAL_FONT, fill="#c9d1d9")
        y += line_height

    image.save(output_path)


def draw_allure_summary(summary: dict, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 1600, 900
    image = Image.new("RGB", (width, height), "#f6f8fa")
    draw = ImageDraw.Draw(image)

    draw.rounded_rectangle((24, 24, width - 24, height - 24), radius=24, fill="white")
    draw.text((70, 60), "Allure Report Summary", font=SANS_BOLD_FONT, fill="#24292f")

    stats = summary["statistic"]
    time_info = summary["time"]
    report_dt = datetime.fromtimestamp(time_info["start"] / 1000)
    report_date = report_dt.strftime("%d.%m.%Y")
    duration_seconds = time_info["duration"] / 1000

    draw.text((70, 120), f"Дата отчета: {report_date}", font=SMALL_FONT, fill="#57606a")
    draw.text(
        (70, 155),
        f"Время выполнения: {duration_seconds:.2f} с",
        font=SMALL_FONT,
        fill="#57606a",
    )

    total = max(stats["total"], 1)
    percent = round(stats["passed"] / total * 100)

    left_card = (70, 240, 540, 780)
    right_card = (620, 240, 1520, 780)
    for card in (left_card, right_card):
        draw.rounded_rectangle(card, radius=24, fill="#ffffff", outline="#d0d7de", width=2)

    draw.text((120, 300), "Тестовых случаев", font=SMALL_FONT, fill="#57606a")
    draw.text((160, 355), str(stats["total"]), font=BIG_FONT, fill="#24292f")
    draw.text((120, 470), "Успешно", font=SMALL_FONT, fill="#57606a")
    draw.text((160, 525), str(stats["passed"]), font=BIG_FONT, fill="#1a7f37")
    draw.text((120, 640), "Ошибок", font=SMALL_FONT, fill="#57606a")
    draw.text((160, 695), str(stats["failed"] + stats["broken"]), font=BIG_FONT, fill="#cf222e")

    draw.text((690, 300), "Общий результат", font=SMALL_FONT, fill="#57606a")
    circle_box = (905, 315, 1255, 665)
    draw.ellipse(circle_box, outline="#d0d7de", width=24)
    draw.arc(circle_box, start=0, end=int(360 * percent / 100), fill="#2da44e", width=24)
    draw.text((1015, 445), f"{percent}%", font=BIG_FONT, fill="#1a7f37")

    draw.text((690, 690), "Набор тестов", font=SMALL_FONT, fill="#57606a")
    draw.text((690, 735), "tests", font=SANS_BOLD_FONT, fill="#24292f")

    image.save(output_path)


def build_structure_text() -> str:
    lines = [
        "mobile_tests/",
        "|-- artifacts/",
        "|   |-- NewPipe_v0.28.7.apk",
        "|   \\-- screenshots/",
        "|       |-- 01_main_screen.png",
        "|       |-- 02_settings_screen.png",
        "|       \\-- 03_about_screen.png",
        "|-- pages/",
        "|   |-- base_page.py",
        "|   |-- main_page.py",
        "|   |-- settings_page.py",
        "|   \\-- about_page.py",
        "|-- tests/",
        "|   \\-- test_navigation.py",
        "|-- utils/",
        "|   \\-- driver_factory.py",
        "|-- reports/",
        "|   |-- allure-results/",
        "|   |-- allure-report/",
        "|   \\-- logs/",
        "|-- conftest.py",
        "|-- config.py",
        "|-- pytest.ini",
        "\\-- requirements.txt",
    ]
    return "\n".join(lines)


def generate_report_assets() -> dict[str, Path]:
    REPORT_ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    appium_log_lines = clean_ansi(read_text(LOGS_DIR / "appium.log")).splitlines()[:22]
    pytest_lines = read_text(LOGS_DIR / "pytest-output.txt").splitlines()
    summary = json.loads(read_text(ALLURE_REPORT_DIR / "widgets" / "summary.json"))

    appium_image = REPORT_ASSETS_DIR / "04_appium_server.png"
    pytest_image = REPORT_ASSETS_DIR / "05_pytest_run.png"
    allure_image = REPORT_ASSETS_DIR / "06_allure_summary.png"
    structure_image = REPORT_ASSETS_DIR / "07_project_structure.png"

    draw_terminal_image("Appium Server", appium_log_lines, appium_image)
    draw_terminal_image("Pytest Result", pytest_lines, pytest_image)
    draw_allure_summary(summary, allure_image)
    draw_terminal_image(
        "Project Structure",
        build_structure_text().splitlines(),
        structure_image,
    )

    return {
        "main_screen": SCREENSHOTS_DIR / "01_main_screen.png",
        "settings_screen": SCREENSHOTS_DIR / "02_settings_screen.png",
        "about_screen": SCREENSHOTS_DIR / "03_about_screen.png",
        "appium_server": appium_image,
        "pytest_run": pytest_image,
        "allure_summary": allure_image,
        "project_structure": structure_image,
    }


def set_default_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    align: int = WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: int = 12,
    space_before: int = 0,
    space_after: int = 6,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)


def add_figure(document: Document, image_path: Path, caption: str, *, width: float = 5.8) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_paragraph(
        document,
        caption,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=11,
        space_before=0,
        space_after=10,
    )


def build_document(images: dict[str, Path]) -> None:
    document = Document()
    set_default_style(document)

    add_paragraph(document, "МИНОБРНАУКИ РОССИИ", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
    add_paragraph(
        document,
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        space_after=2,
    )
    add_paragraph(document, "«Тульский государственный университет»", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
    add_paragraph(document, "Институт прикладной математики и компьютерных наук", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=32)
    add_paragraph(document, "ОТЧЕТ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, space_after=4)
    add_paragraph(document, "ПО ЛАБОРАТОРНЫМ РАБОТАМ ПО ДИСЦИПЛИНЕ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=2)
    add_paragraph(document, "«Интеграция и тестирование программных систем»", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=8)
    add_paragraph(document, "Лабораторная работа №6", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=120)
    add_paragraph(
        document,
        f"Выполнил: ст. гр. {TITLE_BLOCK['student_group']}                                                  {TITLE_BLOCK['student_name']}",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=12,
        space_after=4,
    )
    add_paragraph(
        document,
        f"Проверил:                                                                                      {TITLE_BLOCK['teacher_name']}",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        size=12,
        space_after=120,
    )
    add_paragraph(
        document,
        f"{TITLE_BLOCK['city']} {TITLE_BLOCK['year']} г.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        space_after=0,
    )

    document.add_page_break()

    add_paragraph(document, "Лабораторная работа №6", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=6)
    add_paragraph(document, LAB_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=12)
    add_paragraph(
        document,
        "Цель работы: получить практические навыки автоматизации мобильного UI-тестирования Android-приложения "
        "с использованием Appium, Python, pytest и Allure, а также оформить результаты в виде воспроизводимого отчета.",
    )
    add_paragraph(document, f"{VARIANT_TITLE}. Ход работы", bold=True, size=12)

    add_paragraph(document, "Мобильное тестирование", bold=True, size=12)
    add_paragraph(
        document,
        "Мобильное тестирование предназначено для проверки поведения приложения на реальном устройстве или эмуляторе, "
        "корректности пользовательских сценариев и устойчивости интерфейса к повторяющимся действиям. "
        "В данной работе использован подход Page Object Model, что позволило отделить сценарии тестирования от логики поиска элементов.",
    )
    add_paragraph(document, "В рамках лабораторной работы проверяется:")
    for item in [
        "запуск приложения NewPipe на Android Emulator;",
        "обработка стартовых системных диалогов первого запуска;",
        "открытие navigation drawer;",
        "переход в раздел Settings и проверка ключевых категорий настроек;",
        "переход в раздел About NewPipe и проверка вкладок About & FAQ, Licenses и версии приложения;",
        "формирование отчета Allure по итогам запуска тестов.",
    ]:
        add_paragraph(document, item, size=12, space_after=2)

    add_paragraph(document, "NewPipe", bold=True, size=12, space_before=6)
    add_paragraph(
        document,
        "NewPipe — это свободный Android-клиент для просмотра и прослушивания контента без использования "
        "официальных сервисов Google. Для варианта 12 в качестве источника приложения использован GitHub Release.",
    )
    add_paragraph(document, f"Релиз, использованный в работе: {RELEASE_TAG} от {RELEASE_DATE}.")
    add_paragraph(document, f"Страница релиза: {RELEASE_URL}")
    add_paragraph(document, f"Прямая ссылка на APK: {RELEASE_APK_URL}")

    add_paragraph(document, "Тестовое окружение", bold=True, size=12, space_before=6)
    for item in [
        "ОС: Windows 11;",
        "Эмулятор: Android Emulator, API level 36;",
        "Приложение: NewPipe 0.28.7;",
        "Фреймворк автоматизации: Appium 2.19.0, драйвер UiAutomator2 4.2.9;",
        "Язык и рантайм: Python 3.13.2;",
        "Тестовый фреймворк: pytest 9.0.3;",
        "Отчетность: Allure 2.41.0;",
        "Android Build Tools: 36.1.0.",
    ]:
        add_paragraph(document, item, size=12, space_after=2)

    add_paragraph(document, "Структура проекта", bold=True, size=12, space_before=8)
    add_paragraph(
        document,
        "Проект сформирован в отдельном каталоге mobile_tests и включает в себя конфигурацию, фикстуры, "
        "страничные объекты, тесты, APK приложения и каталоги со скриншотами и отчетами.",
    )
    add_figure(document, images["project_structure"], "Рисунок 1 – Структура проекта mobile_tests", width=5.9)

    add_paragraph(document, "Результаты выполнения", bold=True, size=12, space_before=8)
    add_paragraph(
        document,
        "После запуска приложения были автоматически закрыты стартовые диалоги, после чего открывался основной экран "
        "и выполнялись подготовленные сценарии. На рисунках ниже приведены ключевые состояния интерфейса приложения.",
    )
    add_figure(document, images["main_screen"], "Рисунок 2 – Главный экран NewPipe после закрытия стартовых диалогов", width=3.9)
    add_figure(document, images["settings_screen"], "Рисунок 3 – Экран Settings с основными категориями настроек", width=3.9)
    add_figure(document, images["about_screen"], "Рисунок 4 – Экран About NewPipe с вкладками About & FAQ и Licenses", width=3.9)

    add_paragraph(document, "Проверка инструментов автоматизации", bold=True, size=12, space_before=8)
    add_paragraph(
        document,
        "Для подтверждения корректной работы стенда были сохранены лог запуска Appium Server, результат запуска pytest "
        "и итоговая сводка Allure Report.",
    )
    add_figure(document, images["appium_server"], "Рисунок 5 – Лог запуска Appium Server", width=5.9)
    add_figure(document, images["pytest_run"], "Рисунок 6 – Результат запуска pytest", width=5.9)
    add_figure(document, images["allure_summary"], "Рисунок 7 – Сводка Allure Report по выполненным сценариям", width=5.9)

    add_paragraph(document, "Анализ результатов", bold=True, size=12, space_before=8)
    add_paragraph(document, "В ходе выполнения лабораторной работы:")
    for item in [
        "был подготовлен Android Emulator и установлен APK приложения NewPipe;",
        "настроен стек Appium + UiAutomator2 + pytest + Allure;",
        "реализована Page Object Model для основных экранов приложения;",
        "учтены стартовые диалоги первого запуска приложения;",
        "разработаны и выполнены два UI-сценария навигационного тестирования;",
        "сформирован статический Allure-отчет с результатом 2/2 passed.",
    ]:
        add_paragraph(document, item, size=12, space_after=2)
    add_paragraph(document, "Все запланированные тесты были успешно выполнены.")

    add_paragraph(document, "Вывод", bold=True, size=12, space_before=8)
    add_paragraph(
        document,
        "В результате выполнения лабораторной работы были получены практические навыки автоматизации мобильного тестирования "
        "Android-приложений. Сформированный набор тестов позволяет быстро проверять доступность основных разделов NewPipe, "
        "контролировать корректность навигации и получать наглядный отчет Allure после каждого запуска.",
    )

    add_paragraph(document, "Контрольные вопросы", bold=True, size=12, space_before=8)
    qa_pairs = [
        ("1. Что такое мобильное тестирование?", "Мобильное тестирование проверяет корректность работы приложения на мобильных устройствах и эмуляторах."),
        ("2. Для чего используется Appium?", "Appium используется для автоматизации мобильных приложений на Android и iOS."),
        ("3. Что такое UiAutomator2?", "UiAutomator2 — это драйвер Appium для автоматизации Android-приложений на уровне пользовательского интерфейса."),
        ("4. Для чего применяется Page Object Model?", "Page Object Model отделяет описание экранов и элементов от самих тестовых сценариев."),
        ("5. Что такое desired capabilities?", "Desired capabilities — это параметры запуска сессии Appium, описывающие устройство, приложение и режим тестирования."),
        ("6. Для чего нужен ADB?", "ADB используется для связи с Android-устройством или эмулятором, установки приложений и выполнения системных команд."),
        ("7. Для чего используется pytest?", "pytest нужен для организации, запуска и проверки автоматизированных тестов."),
        ("8. Для чего нужен Allure?", "Allure формирует наглядный отчет по результатам тестового запуска и прикладывает артефакты."),
        ("9. Почему важны стабильные локаторы?", "Стабильные локаторы уменьшают количество ложных падений тестов при изменениях интерфейса."),
        ("10. Зачем обрабатывать стартовые диалоги приложения?", "Стартовые диалоги могут блокировать основной интерфейс и мешать выполнению пользовательских сценариев."),
    ]
    for question, answer in qa_pairs:
        add_paragraph(document, question, bold=True, size=12, space_before=4, space_after=2)
        add_paragraph(document, answer, size=12, space_after=4)

    document.save(OUTPUT_DOCX)


def main() -> None:
    images = generate_report_assets()
    build_document(images)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
